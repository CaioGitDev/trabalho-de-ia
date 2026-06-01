"""Gera o relatorio final do projeto em Word (.docx) e o diagrama do pipeline.

Produz:
  - docs/figures/00_pipeline.png   (diagrama do processo)
  - docs/Relatorio_Final.docx      (documento completo, PT-PT AO90)

Execucao:
    pip install python-docx matplotlib
    python src/gerar_relatorio_docx.py

O texto e escrito de forma direta e tecnica, cobrindo todas as fases, a
reutilizacao dos exemplos das aulas, o treino dos modelos e os conceitos
aplicados, terminando com conclusao e aprendizagem.
"""
from __future__ import annotations

from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
FIG_DIR = ROOT / "docs" / "figures"
DIAGRAM = FIG_DIR / "00_pipeline.png"
MODEL_FIG = FIG_DIR / "08_model_comparison.png"
OUT_DOCX = ROOT / "docs" / "Relatorio_Final.docx"


# ----------------------------------------------------------------------
# Diagrama do pipeline
# ----------------------------------------------------------------------

def build_diagram() -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    fig, ax = plt.subplots(figsize=(9.2, 9.6))
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 12)
    ax.axis("off")

    def box(x, y, w, h, text, color, fontsize=10.0):
        ax.add_patch(FancyBboxPatch(
            (x, y), w, h, boxstyle="round,pad=0.08,rounding_size=0.15",
            linewidth=1.2, edgecolor="#33373d", facecolor=color))
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center",
                fontsize=fontsize, wrap=True)

    def arrow(x, y0, y1):
        ax.annotate("", xy=(x, y1), xytext=(x, y0),
                    arrowprops=dict(arrowstyle="-|>", lw=1.6, color="#33373d"))

    cx, w, h = 3.3, 4.7, 1.05
    steps = [
        ("Imagem do rotulo (fotografia)", "#eef2f7"),
        ("FASE 3 - OCR\npre-processamento (OpenCV) -> Tesseract (por+eng)\n-> limpeza e extracao dos ingredientes", "#dbe7ff"),
        ("FASE 4 - NLP\nnormalizacao + tokenizacao + matching difuso\n(Levenshtein / Jaccard) vs lexico PT/EN/FR", "#dbe7ff"),
        ("FASE 5 - CLASSIFICACAO\nregras especializadas + 3 modelos:\nNaive Bayes, Arvore de Decisao, Random Forest", "#ffe9cc"),
        ("FASE 6 - EXPLICABILIDADE\ntabela por ingrediente (X / OK / ?) + justificacao", "#e6f5e6"),
        ("FASE 7 - RELATORIO / APLICACAO\nRESUMO . ANALISE . ALERGENIOS . RECOMENDACOES", "#e6f5e6"),
    ]
    n = len(steps)
    top = 11.2
    gap = 1.72
    ys = [top - i * gap for i in range(n)]
    for (text, color), y in zip(steps, ys):
        box(cx, y - h, w, h, text, color)
    for i in range(n - 1):
        arrow(cx + w / 2, ys[i] - h, ys[i + 1])

    # caixa lateral: dados + base de conhecimento (Fase 1-2)
    side_y = ys[2] - h - 0.15
    box(0.1, side_y, 2.55, 1.75,
        "FASES 1-2\nOpen Food Facts\n(21 262 produtos PT)\n+ base de\nconhecimento\n(gluten / ambiguos /\n14 alergenios UE)",
        "#f3e8ff", fontsize=9.0)
    # setas da caixa lateral para NLP e Classificacao
    ax.annotate("", xy=(cx, ys[2] - h / 2), xytext=(2.7, side_y + 1.3),
                arrowprops=dict(arrowstyle="-|>", lw=1.3, color="#7a4fb5", linestyle="--"))
    ax.annotate("", xy=(cx, ys[3] - h / 2), xytext=(2.7, side_y + 0.4),
                arrowprops=dict(arrowstyle="-|>", lw=1.3, color="#7a4fb5", linestyle="--"))

    # avaliacao (Fase 8) ao lado do relatorio
    box(8.55, ys[4] - h, 1.65, 2.4,
        "FASE 8\nAVALIACAO\nAccuracy\nPrecision\nRecall / F1\n(matriz de\nconfusao)",
        "#ffe0e0", fontsize=9.0)
    ax.annotate("", xy=(8.55, ys[4] - h / 2), xytext=(cx + w, ys[4] - h / 2),
                arrowprops=dict(arrowstyle="-|>", lw=1.3, color="#b53f3f", linestyle="--"))

    ax.set_title("Arquitetura do sistema — pipeline em 8 fases", fontsize=13, pad=8)
    fig.savefig(DIAGRAM, dpi=130, bbox_inches="tight")
    plt.close(fig)


# ----------------------------------------------------------------------
# Helpers de documento
# ----------------------------------------------------------------------

def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9.5)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)
    return table


def p(doc, text, size=11):
    par = doc.add_paragraph()
    run = par.add_run(text)
    run.font.size = Pt(size)
    return par


def build_document() -> None:
    doc = Document()

    # estilo base
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # ---------------- Capa ----------------
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("Deteção de Glúten e Alergénios\nem Rótulos Alimentares")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1f, 0x3b, 0x73)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Relatório Final do Projeto")
    r.font.size = Pt(15)
    r.italic = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        "Introdução à Inteligência Artificial\n"
        "Universidade Fernando Pessoa — 2026\n\n"
        "Caio Rosa (160173003)\n"
        "Lucélia Ferreira (240001227)\n"
        "Nadine Jeremias (240001501)"
    ).font.size = Pt(12)

    doc.add_page_break()

    # ---------------- Resumo ----------------
    doc.add_heading("Resumo", level=1)
    p(doc,
      "Este projeto desenvolve um sistema de Inteligência Artificial que analisa "
      "a fotografia de um rótulo alimentar e identifica automaticamente a presença "
      "de glúten e de alergénios, assinala ingredientes ambíguos e atribui um grau "
      "de risco, sempre com uma justificação por ingrediente. A solução combina "
      "reconhecimento ótico de carateres (OCR), processamento de linguagem natural "
      "(PLN), um classificador baseado em regras e modelos de aprendizagem "
      "supervisionada (Naive Bayes, Árvore de Decisão e Random Forest). Todas as "
      "etapas reutilizam conceitos lecionados na unidade curricular. Na avaliação "
      "não circular sobre texto real, o sistema atinge 91,2 % de exatidão na "
      "deteção de glúten; o melhor modelo supervisionado (Random Forest) alcança "
      "88,7 % de exatidão no conjunto de teste.")

    # ---------------- 1. Introdução ----------------
    doc.add_heading("1. Introdução e objetivo", level=1)
    p(doc,
      "A doença celíaca e as alergias alimentares obrigam muitos consumidores a "
      "ler cuidadosamente os rótulos. O Regulamento (UE) n.º 1169/2011 torna "
      "obrigatória a declaração de 14 alergénios (entre os quais os cereais com "
      "glúten, o leite, os ovos, a soja, o amendoim e os frutos de casca rija). "
      "Contudo, a leitura manual é morosa e propensa a erros, e a origem de "
      "alguns ingredientes (amido, aromas, proteína vegetal) nem sempre é clara.")
    p(doc,
      "O objetivo é construir uma aplicação que, a partir de uma imagem de um "
      "rótulo, produza uma análise clara, explicável e justificável: deteção de "
      "glúten (em cinco graus, de «Sem Glúten» a «Contém Glúten»), estado de "
      "alergénios, identificação de ingredientes ambíguos, grau de risco e "
      "recomendações ao consumidor.")

    # ---------------- 2. Reavaliação técnica ----------------
    doc.add_heading("2. Reavaliação técnica da proposta", level=1)
    p(doc,
      "A proposta inicial sugeria «construir um modelo de linguagem (LLM) de "
      "raiz». Após análise crítica, esta via foi rejeitada: treinar um LLM exige "
      "enormes volumes de dados, tempo de treino e recursos computacionais "
      "incompatíveis com o contexto académico, e seria desadequado a um problema "
      "que é, na sua essência, de extração e classificação de informação "
      "estruturada.")
    p(doc,
      "Optámos por uma arquitetura modular — OCR + PLN + regras especializadas + "
      "classificadores supervisionados — mais transparente, mais barata de treinar "
      "e diretamente alinhada com a matéria lecionada. Esta escolha privilegia "
      "também a explicabilidade, requisito central do enunciado: ao contrário de "
      "um modelo «caixa-negra», cada decisão é justificada ao nível do ingrediente.")

    # ---------------- 3. Arquitetura ----------------
    doc.add_heading("3. Arquitetura geral do sistema", level=1)
    p(doc,
      "O sistema está organizado em oito fases. As fases 1 e 2 (investigação e "
      "construção do conjunto de dados) sustentam o resto: definem a base de "
      "conhecimento e o conjunto de treino. As fases 3 a 7 formam o fluxo de "
      "produção — da imagem ao relatório — e a fase 8 mede o desempenho. O diagrama "
      "seguinte resume todo o processo.")
    if DIAGRAM.exists():
        doc.add_picture(str(DIAGRAM), width=Inches(5.6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p(doc,
      "Cada fase comunica com a seguinte através de estruturas de dados bem "
      "definidas (o resultado da análise de PLN alimenta as regras, que alimentam "
      "o motor de explicabilidade, que alimenta o relatório), o que torna o "
      "código modular e testável de forma independente.", size=10)

    # ---------------- 4. Dados ----------------
    doc.add_heading("4. Dados e base de conhecimento (Fases 1–2)", level=1)
    p(doc,
      "A fonte de dados é o Open Food Facts, base alimentar pública com forte "
      "presença de produtos europeus. A partir do dump oficial (~9 GiB, lido em "
      "modo de fluxo para não esgotar a memória) extraímos 21 262 produtos "
      "comercializados em Portugal, guardando os campos relevantes: nome, marca, "
      "categorias, ingredientes, alergénios, traços, etiquetas, Nutri-Score e país.")
    p(doc,
      "A preparação dos dados incluiu limpeza (remoção de registos sem "
      "ingredientes e de duplicados), normalização de texto e uniformização. "
      "Verificámos que o texto livre do dump contém «mojibake» (carateres "
      "corrompidos), pelo que demos prioridade à coluna normalizada de tags de "
      "ingredientes (taxonomia «en:») para a análise estatística e para o treino.")
    p(doc,
      "A base de conhecimento (módulo knowledge_base) codifica o saber do domínio: "
      "as tags de cereais com glúten, as tags de ingredientes ambíguos e os 14 "
      "alergénios obrigatórios da União Europeia, com a respetiva designação em "
      "português. É o núcleo simbólico que as fases seguintes consultam.")

    # ---------------- 5. OCR ----------------
    doc.add_heading("5. Extração de texto — OCR (Fase 3)", level=1)
    p(doc,
      "O módulo de OCR converte a imagem em texto. O fluxo é: pré-processamento "
      "com OpenCV (escala de cinzentos, redução de ruído, binarização e correção "
      "de inclinação), reconhecimento com o motor Tesseract configurado para "
      "português e inglês, limpeza do texto (correção de erros típicos, junção de "
      "palavras hifenizadas, remoção de ruído de margens) e, por fim, isolamento "
      "da secção «Ingredientes» por heurística. A escolha do Tesseract justifica-se "
      "por ser livre, maduro e com bom suporte de português.")

    # ---------------- 6. NLP ----------------
    doc.add_heading("6. Processamento de linguagem natural (Fase 4)", level=1)
    p(doc,
      "Esta fase transforma texto livre numa análise estruturada. As etapas são a "
      "normalização (forma Unicode, minúsculas, remoção de acentos para efeitos de "
      "comparação), a tokenização da lista de ingredientes (separação por vírgulas, "
      "parênteses, percentagens e conjunções «e/y/and») e a remoção de "
      "inconsistências (números, códigos de aditivos, ruído).")
    p(doc,
      "O coração da fase é a correspondência (matching) entre cada segmento de "
      "texto e um léxico de ingredientes em português, inglês e francês, que faz a "
      "ponte para as tags «en:» da base de conhecimento. A correspondência é "
      "primeiro exata e, quando falha — algo frequente devido a erros de OCR como "
      "«Batato» por «batata» ou «centeo» por «centeio» — recorre a métricas de "
      "distância. Aqui aplicámos diretamente os conceitos das aulas de Estilometria: "
      "a distância de Levenshtein (com um rácio de semelhança normalizado) e o "
      "índice de Jaccard, ambos implementados em Python puro. A correspondência é "
      "ordenada por especificidade, para que «farinha de trigo» (glúten) prevaleça "
      "sobre «farinha» (ambíguo) e «amido de milho» (seguro) sobre «amido» "
      "(ambíguo). O sistema deteta ainda as declarações «contém», «pode conter "
      "traços» e o rótulo «sem glúten».")

    # ---------------- 7. Classificação ----------------
    doc.add_heading("7. Classificação e treino dos modelos (Fase 5)", level=1)
    p(doc,
      "O enunciado pede a comparação de várias abordagens. Implementámos primeiro "
      "um classificador baseado em regras especializadas, que a partir das tags de "
      "um produto atribui o estado de glúten («sem», «contém» ou «suspeito») e que "
      "constitui a referência explicável do sistema.")
    p(doc,
      "Em seguida treinámos modelos de aprendizagem supervisionada. Como não "
      "dispúnhamos de rótulos humanos para todos os produtos, usámos uma estratégia "
      "de supervisão fraca: as próprias regras geram o rótulo de cada produto, "
      "funcionando como «professor» automático. Cada produto é representado por um "
      "vetor binário do tipo «saco de tags» (bag-of-tags): para cada tag de "
      "ingrediente do vocabulário, 1 se está presente, 0 caso contrário. O conjunto "
      "resultante tem 8 770 amostras, 968 atributos e 3 classes.")
    p(doc,
      "Comparámos três algoritmos lecionados: Naive Bayes de Bernoulli, "
      "implementado de raiz a partir do Teorema de Bayes (verosimilhança de "
      "Bernoulli com suavização de Laplace e cálculo em espaço logarítmico para "
      "evitar underflow); a Árvore de Decisão; e a Random Forest. O protocolo de "
      "treino seguiu o que foi ensinado: divisão em treino e teste (train_test_split) "
      "estratificada, validação cruzada K-Fold estratificada com 5 folds, e o ciclo "
      "fit/predict. A avaliação reutiliza o código do professor (confusion_matrix2), "
      "que calcula a matriz de confusão e as métricas Accuracy, Precision, Recall e "
      "F1. A nossa implementação de Naive Bayes foi validada por coincidir 100 % "
      "com a do scikit-learn.")
    add_table(doc,
              ["Modelo", "K-Fold (exatidão)", "Teste (exatidão)", "Macro-F1"],
              [["Naive Bayes (de raiz)", "81,3 %", "82,0 %", "80,8 %"],
               ["Árvore de Decisão", "87,8 %", "87,9 %", "87,0 %"],
               ["Random Forest", "89,5 %", "88,7 %", "88,1 %"]])
    if MODEL_FIG.exists():
        doc.add_picture(str(MODEL_FIG), width=Inches(5.4))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p(doc,
      "A Random Forest obteve o melhor desempenho. O Naive Bayes ficou abaixo "
      "sobretudo na classe «suspeito» (a mais difícil para todos os modelos), o "
      "que é coerente com a sua hipótese de independência entre atributos. Note-se "
      "que o problema não é trivial: o vocabulário exclui tags muito raras e o "
      "rótulo depende também de campos que não entram nos atributos, pelo que existe "
      "aprendizagem real e não mera memorização das regras.", size=10)

    # ---------------- 8. Explicabilidade ----------------
    doc.add_heading("8. Explicabilidade (Fase 6)", level=1)
    p(doc,
      "O sistema justifica todas as decisões. Para cada ingrediente reconhecido "
      "produz uma linha com um símbolo de estado — ❌ contém glúten, ⚠️ ambíguo ou "
      "alergénio, ✅ seguro — e uma observação textual, incluindo a tag «en:» de "
      "origem (rastreabilidade) e a confiança quando a correspondência foi "
      "aproximada. Esta camada não acrescenta lógica de decisão: torna transparente "
      "aquilo que as regras e o PLN já determinaram.")

    # ---------------- 9. Relatório/App ----------------
    doc.add_heading("9. Relatório e aplicação (Fase 7)", level=1)
    p(doc,
      "Todo o pipeline culmina num relatório automático com quatro secções: "
      "RESUMO (classificação global), ANÁLISE DOS INGREDIENTES (tabela detalhada), "
      "ALERGÉNIOS DETETADOS e RECOMENDAÇÕES. As recomendações são geradas por "
      "regras a partir do grau de risco, dos ingredientes ambíguos, dos alergénios "
      "e da qualidade do reconhecimento (por exemplo, «produto inadequado para "
      "celíacos» ou «confirmar a origem do amido junto do fabricante»). "
      "Disponibilizámos o sistema como linha de comandos e como aplicação web "
      "(Flask), onde o utilizador cola o texto ou carrega a foto do rótulo e recebe "
      "o relatório com cores por grau de risco.")

    # ---------------- 10. Avaliação ----------------
    doc.add_heading("10. Avaliação e validação (Fase 8)", level=1)
    p(doc,
      "Avaliámos o sistema em três frentes, sempre com as métricas e a matriz de "
      "confusão do professor.")
    p(doc,
      "A) Modelo supervisionado: a Random Forest obteve 88,7 % de exatidão e "
      "88,1 % de macro-F1 no conjunto de teste retido.")
    p(doc,
      "B) Pipeline completo, de forma não circular: corremos o PLN e as regras "
      "sobre o texto livre de ingredientes e comparámos a deteção de glúten com um "
      "rótulo independente — o campo de alergénios do Open Food Facts, que o "
      "pipeline nunca lê. Em 1 500 produtos obtivemos 91,2 % de exatidão, com "
      "precisão de 85,7 % e recall de 89,9 % na deteção de glúten.")
    add_table(doc,
              ["Avaliação", "Exatidão", "Precisão (glúten)", "Recall (glúten)"],
              [["Random Forest (teste)", "88,7 %", "—", "—"],
               ["Pipeline vs rótulo independente", "91,2 %", "85,7 %", "89,9 %"]])
    p(doc,
      "C) Impacto do OCR: nas fotografias reais (confiança média ~48 %) o veredicto "
      "obtido a partir da imagem é tipicamente menos grave do que o obtido a partir "
      "do texto limpo, evidenciando que a qualidade fotográfica é a principal "
      "fragilidade.")
    p(doc,
      "A análise de erros foi reveladora. Antes de alargar o léxico, os falsos "
      "negativos deviam-se sobretudo a rótulos em inglês e francês; juntar esses "
      "termos elevou a exatidão da avaliação B de ~88 % para 91,2 % e o recall de "
      "glúten de 80 % para ~90 %. Vários «falsos positivos» revelaram-se, na "
      "verdade, deteções corretas que o rótulo do Open Food Facts tinha omitido, "
      "pelo que a precisão real é superior à medida.")

    # ---------------- 11. Reutilização ----------------
    doc.add_heading("11. Reutilização dos conceitos das aulas", level=1)
    p(doc,
      "Um requisito central foi demonstrar a aplicação prática da matéria. A tabela "
      "seguinte resume onde cada conceito foi usado.")
    add_table(doc,
              ["Conceito lecionado", "Onde foi aplicado"],
              [["Teorema de Bayes / Naive Bayes", "Classificador implementado de raiz"],
               ["Distâncias Levenshtein, Jaccard, Hamming", "Correspondência difusa de texto de OCR"],
               ["train_test_split, fit, predict", "Treino dos classificadores"],
               ["Validação cruzada K-Fold estratificada", "Estimativa robusta de desempenho"],
               ["Árvore de Decisão e Random Forest", "Modelos comparados (scikit-learn)"],
               ["Matriz de confusão e métricas", "Avaliação (código confusion_matrix2 do professor)"]])
    p(doc,
      "O código de matriz de confusão fornecido nas aulas foi reutilizado "
      "literalmente, tanto na comparação de modelos (Fase 5) como na validação "
      "final (Fase 8), garantindo coerência entre o que foi ensinado e o que foi "
      "medido.", size=10)

    # ---------------- 12. Limitações ----------------
    doc.add_heading("12. Limitações", level=1)
    for item in [
        "O OCR é a maior fonte de erro: em fotografias de baixa qualidade parte do "
        "texto não é recuperada e o veredicto tende a subestimar o risco.",
        "O léxico cobre português, inglês e francês; outras línguas (italiano, "
        "espanhol) ainda geram falsos negativos.",
        "A correspondência difusa favorece o recall (mais seguro sobre-assinalar "
        "risco), o que introduz alguns falsos positivos.",
        "Os classificadores usam rótulos gerados pelas próprias regras (supervisão "
        "fraca): medem a capacidade de reproduzir as regras, não um juízo humano "
        "independente.",
        "O rótulo de alergénios usado como referência na Fase 8 é preenchido por "
        "contribuidores e tem omissões, o que adiciona ruído à medição.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # ---------------- 13. Conclusão ----------------
    doc.add_heading("13. Conclusão — considerações finais e aprendizagem", level=1)
    p(doc,
      "Cumprimos todos os objetivos do enunciado: construímos uma aplicação que "
      "recebe a fotografia de um rótulo e produz uma análise automática, explicável "
      "e justificada do risco de glúten e de alergénios, com resultados sólidos "
      "(91,2 % de exatidão na deteção de glúten sobre texto real).")
    p(doc,
      "A principal aprendizagem foi confirmar que, para um problema bem delimitado, "
      "uma arquitetura clássica e modular pode ser mais adequada — e muito mais "
      "transparente — do que uma solução monolítica de grande dimensão. Cada peça "
      "(OCR, PLN, regras, classificadores) resolve uma parte do problema e pode ser "
      "validada isoladamente, o que facilitou o desenvolvimento e a deteção de "
      "erros.")
    p(doc,
      "Aprendemos também a importância de uma avaliação honesta. Ao medir o "
      "pipeline contra um rótulo independente, em vez dos rótulos gerados pelas "
      "nossas próprias regras, evitámos uma avaliação circular e expusemos as "
      "verdadeiras fragilidades — desde logo a cobertura linguística do léxico, que "
      "corrigimos e cuja melhoria pudémos quantificar. A própria análise de erros "
      "ensinou-nos a distinguir falhas do sistema de omissões dos dados de origem.")
    p(doc,
      "Do ponto de vista dos conteúdos, o projeto consolidou a compreensão do "
      "Teorema de Bayes (ao implementar o Naive Bayes de raiz), das métricas de "
      "distância em PLN, do protocolo de treino e validação (train/test, K-Fold) e "
      "da leitura de uma matriz de confusão. Por fim, num domínio ligado à saúde, "
      "interiorizámos que a explicabilidade não é um extra mas um requisito: um "
      "sistema que aconselha sobre alimentos tem de justificar cada decisão para "
      "ser digno de confiança.")
    p(doc,
      "Como trabalho futuro, identificamos o reforço do OCR, o alargamento do "
      "léxico a mais línguas e a recolha de um pequeno conjunto de rótulos com "
      "verdade anotada por humanos, que permitiria uma avaliação ainda mais "
      "rigorosa.")

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)


def main() -> None:
    build_diagram()
    build_document()
    print(f"Diagrama: {DIAGRAM.relative_to(ROOT)}")
    print(f"Relatorio: {OUT_DOCX.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
